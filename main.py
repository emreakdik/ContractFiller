import os
import docx


def create_contract(contract_template, data):
    doc = docx.Document(contract_template)

    for para in doc.paragraphs:
        for key in data.keys():
            para.text = para.text.replace("{" + key + "}", data[key])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key in data.keys():
                    cell.text = cell.text.replace("{" + key + "}", data[key])

    return doc


templates_dir = "templates"

templates = []
for filename in os.listdir(templates_dir):
    if filename.endswith(".docx"):
        templates.append(os.path.join(templates_dir, filename))

directory = "Filled Contracts"
if not os.path.exists(directory):
    os.makedirs(directory)

name = input("Name-Surname: ")
phone_number = input("Phone Number: ")
ID = input("ID: ")
date_of_birth = input("Date of Birth: ")
address = input("Adress: ")
job_adress = input("Occupation Adress: ")
occupation = input("Occupation: ")
start_date = input("Start Date: ")
monthly_wage = input("Wage: ")


data = {"Name": name, "ID": ID, "Date of Birth": date_of_birth, "Address": address,
        "Occupation": occupation, "Start Date": start_date, "Monthly Wage": monthly_wage, "Phone Number": phone_number, "Job Adress": job_adress}

for template in templates:
    doc = create_contract(template, data)
    filename = os.path.splitext(os.path.basename(template))[
        0] + "_" + name + ".docx"
    filepath = os.path.join(directory, filename)
    doc.save(filepath)
